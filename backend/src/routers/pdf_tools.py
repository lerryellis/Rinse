from fastapi import APIRouter, UploadFile, File, HTTPException, Request
from fastapi.responses import StreamingResponse
import io
import fitz  # PyMuPDF
from src.services.usage import check_limits, record_usage
from src.services.auth import require_user

router = APIRouter()


async def enforce_limits(request: Request, file_size: int, tool: str):
    """Require auth, check usage limits with device fingerprint, and record usage."""
    user_id = require_user(request)
    device_id = request.headers.get("x-device-id")
    error = check_limits(user_id, file_size, device_id)
    if error:
        raise HTTPException(status_code=429, detail=error)
    record_usage(user_id, tool, file_size, device_id=device_id)


@router.post("/compress")
async def compress_pdf(request: Request, file: UploadFile = File(...)):
    """Compress a PDF to reduce file size."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    contents = await file.read()
    await enforce_limits(request, len(contents), "compress")
    original_size = len(contents)

    doc = fitz.open(stream=contents, filetype="pdf")

    # Re-save with garbage collection and deflation
    output = doc.tobytes(
        garbage=4,
        deflate=True,
        clean=True,
    )
    doc.close()

    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="compressed.pdf"',
            "X-Original-Size": str(original_size),
            "X-Compressed-Size": str(len(output)),
        },
    )


@router.post("/to-jpg")
async def pdf_to_jpg(request: Request, file: UploadFile = File(...)):
    """Convert PDF pages to JPG images. Returns a zip of images."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    import zipfile

    contents = await file.read()
    await enforce_limits(request, len(contents), "to-jpg")
    doc = fitz.open(stream=contents, filetype="pdf")

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("jpeg")
            zf.writestr(f"page-{i + 1}.jpg", img_bytes)

    doc.close()
    zip_buffer.seek(0)

    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="pdf-pages.zip"',
        },
    )


@router.post("/to-word")
async def pdf_to_word(request: Request, file: UploadFile = File(...)):
    """Convert PDF to Word (.docx)."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    from docx import Document
    from docx.shared import Pt

    contents = await file.read()
    await enforce_limits(request, len(contents), "to-word")
    doc = fitz.open(stream=contents, filetype="pdf")
    word_doc = Document()

    for page in doc:
        blocks = page.get_text("blocks")
        for block in blocks:
            if block[6] == 0:  # text block
                text = block[4].strip()
                if text:
                    p = word_doc.add_paragraph(text)
                    p.style.font.size = Pt(11)
        word_doc.add_page_break()

    doc.close()

    output = io.BytesIO()
    word_doc.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="converted.docx"',
        },
    )


@router.post("/to-excel")
async def pdf_to_excel(request: Request, file: UploadFile = File(...)):
    """Convert PDF tables to Excel (.xlsx)."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    from openpyxl import Workbook

    contents = await file.read()
    await enforce_limits(request, len(contents), "to-excel")
    doc = fitz.open(stream=contents, filetype="pdf")
    wb = Workbook()
    ws = wb.active
    ws.title = "PDF Content"

    row_num = 1
    for page in doc:
        tables = page.find_tables()
        for table in tables:
            for row in table.extract():
                for col_idx, cell in enumerate(row):
                    ws.cell(row=row_num, column=col_idx + 1, value=cell or "")
                row_num += 1
            row_num += 1  # blank row between tables

    doc.close()

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f'attachment; filename="converted.xlsx"',
        },
    )


@router.post("/to-text")
async def pdf_to_text(request: Request, file: UploadFile = File(...)):
    """Extract all text from a PDF."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    contents = await file.read()
    await enforce_limits(request, len(contents), "to-text")
    doc = fitz.open(stream=contents, filetype="pdf")

    text = ""
    for page in doc:
        text += page.get_text() + "\n\n"

    doc.close()

    return StreamingResponse(
        io.BytesIO(text.encode("utf-8")),
        media_type="text/plain",
        headers={
            "Content-Disposition": f'attachment; filename="extracted.txt"',
        },
    )


@router.post("/watermark")
async def add_watermark(request: Request, file: UploadFile = File(...), text: str = "DRAFT"):
    """Add a text watermark to a PDF."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    contents = await file.read()
    await enforce_limits(request, len(contents), "watermark")
    doc = fitz.open(stream=contents, filetype="pdf")

    for page in doc:
        rect = page.rect
        tw = fitz.TextWriter(page.rect)
        font = fitz.Font("helv")
        fontsize = rect.width / len(text) * 1.5
        tw.append(
            fitz.Point(rect.width / 4, rect.height / 2),
            text,
            font=font,
            fontsize=fontsize,
        )
        tw.write_text(page, opacity=0.3, color=(0.7, 0.7, 0.7), rotate=45)

    output = doc.tobytes(garbage=4, deflate=True)
    doc.close()

    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="watermarked.pdf"',
        },
    )


@router.post("/protect")
async def protect_pdf(request: Request, file: UploadFile = File(...), password: str = ""):
    """Password-protect a PDF."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")
    if not password:
        raise HTTPException(status_code=400, detail="Password is required")

    contents = await file.read()
    await enforce_limits(request, len(contents), "protect")
    doc = fitz.open(stream=contents, filetype="pdf")

    output = doc.tobytes(
        encryption=fitz.PDF_ENCRYPT_AES_256,
        user_pw=password,
        owner_pw=password,
        garbage=4,
        deflate=True,
    )
    doc.close()

    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="protected.pdf"',
        },
    )


@router.post("/unlock")
async def unlock_pdf(request: Request, file: UploadFile = File(...), password: str = ""):
    """Remove password from a PDF."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    contents = await file.read()
    await enforce_limits(request, len(contents), "unlock")
    doc = fitz.open(stream=contents, filetype="pdf")

    if doc.is_encrypted:
        if not doc.authenticate(password):
            doc.close()
            raise HTTPException(status_code=400, detail="Incorrect password")

    output = doc.tobytes(garbage=4, deflate=True)
    doc.close()

    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="unlocked.pdf"',
        },
    )


@router.post("/crop")
async def crop_pdf(request: Request, file: UploadFile = File(...)):
    """Auto-crop PDF page margins by detecting content bounds."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    contents = await file.read()
    await enforce_limits(request, len(contents), "crop")
    doc = fitz.open(stream=contents, filetype="pdf")

    for page in doc:
        # Get bounding box of all content on the page
        blocks = page.get_text("blocks")
        if not blocks:
            continue
        x0 = min(b[0] for b in blocks)
        y0 = min(b[1] for b in blocks)
        x1 = max(b[2] for b in blocks)
        y1 = max(b[3] for b in blocks)
        # Add small margin
        margin = 20
        crop = fitz.Rect(
            max(0, x0 - margin),
            max(0, y0 - margin),
            min(page.rect.width, x1 + margin),
            min(page.rect.height, y1 + margin),
        )
        page.set_cropbox(crop)

    output = doc.tobytes(garbage=4, deflate=True)
    doc.close()

    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="cropped.pdf"',
        },
    )


@router.post("/word-to-pdf")
async def word_to_pdf(request: Request, file: UploadFile = File(...)):
    """Convert Word document to PDF."""
    from docx import Document

    contents = await file.read()
    await enforce_limits(request, len(contents), "word-to-pdf")
    word_doc = Document(io.BytesIO(contents))

    # Create a PDF from the text content
    doc = fitz.open()
    for para in word_doc.paragraphs:
        if not para.text.strip():
            continue
        # Simple approach: add text to pages
        if len(doc) == 0:
            page = doc.new_page()
            y_pos = 72
        text = para.text
        page = doc[-1]
        tw = fitz.TextWriter(page.rect)
        font = fitz.Font("helv")
        fontsize = 11
        tw.append(fitz.Point(72, y_pos), text, font=font, fontsize=fontsize)
        tw.write_text(page)
        y_pos += fontsize * 1.5
        if y_pos > page.rect.height - 72:
            page = doc.new_page()
            y_pos = 72

    output = doc.tobytes(garbage=4, deflate=True)
    doc.close()

    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="converted.pdf"',
        },
    )


@router.post("/html-to-pdf")
async def html_to_pdf(request: Request):
    """Convert HTML content or a URL to PDF using a headless browser."""
    from playwright.async_api import async_playwright

    content_type = request.headers.get("content-type", "")

    # Accept JSON body with html or url
    if "application/json" in content_type:
        body = await request.json()
        html = body.get("html")
        url = body.get("url")
        page_format = body.get("format", "A4")
        landscape = body.get("landscape", False)
    # Accept multipart form with an HTML file
    else:
        form = await request.form()
        file = form.get("file")
        url = form.get("url")
        page_format = form.get("format", "A4")
        landscape = str(form.get("landscape", "false")).lower() == "true"
        html = None
        if file and hasattr(file, "read"):
            html_bytes = await file.read()
            html = html_bytes.decode("utf-8", errors="replace")

    if not html and not url:
        raise HTTPException(status_code=400, detail="Provide either 'html' content or a 'url'")

    # Estimate size for limits (use html length or a default for URL)
    size_estimate = len(html.encode("utf-8")) if html else 1024
    user_id = require_user(request)
    device_id = request.headers.get("x-device-id")
    error = check_limits(user_id, size_estimate, device_id)
    if error:
        raise HTTPException(status_code=429, detail=error)

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()

        if url:
            await page.goto(url, wait_until="networkidle", timeout=30000)
        else:
            await page.set_content(html, wait_until="networkidle", timeout=15000)

        pdf_bytes = await page.pdf(
            format=page_format,
            landscape=landscape,
            print_background=True,
            margin={"top": "20mm", "bottom": "20mm", "left": "15mm", "right": "15mm"},
        )

        await browser.close()

    record_usage(user_id, "html-to-pdf", len(pdf_bytes), device_id=device_id)

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": 'attachment; filename="html-converted.pdf"',
        },
    )
